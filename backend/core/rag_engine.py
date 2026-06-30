import logging
import os
import re
import asyncio
from pathlib import Path
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
from config import config

logger = logging.getLogger(__name__)

class RAGEngine:
    def __init__(self):
        self.persist_directory = config.VECTOR_DB_DIR
        self.collection_name = "sales_knowledge"
        self._embeddings = None
        self._vector_store = None
        self._reranker = None

        # Larger chunks with more overlap for better context preservation
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    @property
    def embeddings(self):
        if self._embeddings is None:
            logger.info("Loading embedding model (first use)...")
            self._embeddings = HuggingFaceEmbeddings(
                model_name="BAAI/bge-base-en-v1.5",
                model_kwargs={"device": "cpu"},
                encode_kwargs={"normalize_embeddings": True},
            )
        return self._embeddings

    @property
    def vector_store(self):
        if self._vector_store is None:
            self._vector_store = Chroma(
                collection_name=self.collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.persist_directory,
            )
        return self._vector_store

    @property
    def reranker(self):
        if self._reranker is None:
            self._reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
        return self._reranker

    def process_document(self, file_path, module_id=None, document_id=None):
        """Extracts text from PDF/TXT/MD, splits it, and adds to ChromaDB."""
        try:
            documents = self._load_documents(file_path)
            chunks = self.text_splitter.split_documents(documents)
            if not chunks:
                return False, "No text could be extracted from the document.", 0

            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "module_id": module_id or "default",
                    "document_id": document_id or Path(file_path).stem,
                    "source": file_path,
                    "chunk_index": i,
                    "filename": Path(file_path).name,
                })

            self.vector_store.add_documents(chunks)
            return True, f"Successfully processed and embedded {len(chunks)} chunks.", len(chunks)
        except Exception as e:
            logger.error("Error processing document: %s", e)
            return False, str(e), 0

    async def retrieve_context(self, query, top_k=3, module_id=None, rerank=True):
        """Hybrid retrieval: over-fetch with semantic search, then rerank for precision."""
        try:
            filter_kwargs = {"module_id": module_id} if (module_id and module_id != "ALL") else None
            # Over-fetch 3x candidates for reranking
            fetch_k = top_k * 3 if rerank else top_k
            results = await asyncio.to_thread(
                self.vector_store.similarity_search, query, k=fetch_k, filter=filter_kwargs
            )

            if not results:
                return ""

            if rerank and len(results) > top_k:
                # Keyword boost: prioritize chunks containing query terms
                query_terms = set(re.findall(r'\w+', query.lower()))
                for doc in results:
                    doc_terms = set(re.findall(r'\w+', doc.page_content.lower()))
                    doc.metadata["keyword_overlap"] = len(query_terms & doc_terms)

                # Cross-encoder reranking for final precision
                pairs = [[query, doc.page_content] for doc in results]
                scores = await asyncio.to_thread(self.reranker.predict, pairs)
                scored = sorted(zip(scores, results), key=lambda x: x[0], reverse=True)
                results = [doc for _, doc in scored[:top_k]]
            else:
                results = results[:top_k]

            context = "\n\n---\n\n".join([doc.page_content for doc in results])
            return context
        except Exception as e:
            logger.error("Error retrieving context: %s", e)
            return ""

    def get_all_context_for_module(self, module_id, max_chars=20000):
        """Retrieves all document chunks for a module up to a char limit."""
        try:
            if module_id == "ALL":
                matches = self.vector_store.get()
            else:
                matches = self.vector_store.get(where={"module_id": module_id})
            docs = matches.get("documents", [])
            if not docs:
                return ""

            full_text = ""
            for text in docs:
                if len(full_text) + len(text) > max_chars:
                    break
                full_text += text + "\n\n"
            return full_text.strip()
        except Exception as e:
            logger.error("Error retrieving all module context: %s", e)
            return ""

    def delete_document_vectors(self, module_id, document_id):
        try:
            matches = self.vector_store.get(where={"$and": [{"module_id": module_id}, {"document_id": document_id}]})
            ids = matches.get("ids", [])
            if ids:
                self.vector_store.delete(ids=ids)
            return True, len(ids)
        except Exception as e:
            logger.error("Error deleting document vectors: %s", e)
            return False, 0

    def delete_module_vectors(self, module_id):
        try:
            matches = self.vector_store.get(where={"module_id": module_id})
            ids = matches.get("ids", [])
            if ids:
                self.vector_store.delete(ids=ids)
            return True, len(ids)
        except Exception as e:
            logger.error("Error deleting module vectors: %s", e)
            return False, 0

    def _load_documents(self, file_path):
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            return PyMuPDFLoader(file_path).load()
        if suffix in {".txt", ".md"}:
            text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            return [Document(page_content=text, metadata={"source": file_path})]
        if suffix in {".docx", ".doc"}:
            return self._load_docx(file_path)
        raise ValueError("Only PDF, DOCX, TXT, and Markdown files are supported.")

    def _load_docx(self, file_path):
        """Extracts text from a Word document (paragraphs + tables)."""
        try:
            import docx  # python-docx
        except ImportError:
            raise ValueError(
                "DOCX support requires the 'python-docx' package. "
                "Install it with: pip install python-docx"
            )
        document = docx.Document(file_path)
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        # Include table cell text too
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        if not text.strip():
            raise ValueError("No readable text found in the Word document.")
        return [Document(page_content=text, metadata={"source": file_path})]


rag_engine = RAGEngine()
